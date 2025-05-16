from docx import Document 
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.shared import OxmlElement, qn
import networkx as nx
from telnetlib import Telnet 
import re
import requests
import time

#region variables

# 宣告變數初始化
filename = ""
ip_address = ""
ipList = ""
lines = ""
lldp1 = ""
lldp2 = ""
lldp3 = ""
matches1 = ""
matches2 = ""
matches3 = ""
output = ""
output_str1 = ""
output_str2 = ""
pattern1 = ""
pattern2 = ""
pattern3 = ""
trunk = ""
url_prefix = "https://api.maclookup.app/v2/macs/"

HOST = ""
USER = "admin"
PASSWORDS = ["admin", "1234"]
ZYXEL = 0

#endregion

#region function

# fun1.處理字串
def string_process(i):

    global ip_address, trunk

    # 將trunk預設設為空
    trunk = ""

    # 將每一行字串前後的空白、換行全部清空，如果該行為空字串則直接跳到下一行
    ipList[i] = ipList[i].strip()
    if len(ipList[i]) == 0:
        return 0

    # parse每一行分割成ip跟trunk兩個欄位，如果只有ip的話代表是gateway
    result = re.split(r'\s+', ipList[i])

    if len(result) == 1:
        ip_address = result[0]

    elif len(result) == 2:
        ip_address = result[0]
        trunk = result[1]
    else:
        print("ip.txt每一行的參數過多 => 輸入錯誤!!!")
        return 0
    
    return 1

# fun2.登入並輸入指令
def log_in(PASSWORD):

    global HOST

    # 代入IP
    print("\n=============================\n嘗試登入IP : ", ip_address)
    HOST = ip_address

    # 建立telnet連線(IP)
    tn = Telnet()
    tn.open(HOST)

    # 輸入帳號
    output = tn.read_until(b"Username: ", timeout = 5)
    tn.write(USER.encode("ascii") + b"\n")

    # 由狀態訊息判斷是否登入zyxel
    if output == b'\r\nUser name: ':

        # fun3.登入zyxel
        if log_in_zyxel(PASSWORD, tn) == 0:
            return 0
        
        # 登入成功
        else:
            return 1

    else:
        
        # fun4.登入edgecore
        if log_in_edgecore(PASSWORD, tn) == 0:
            return 0
        
        #登入成功
        else:
            return 1

# fun3.登入zyxel
def log_in_zyxel(PASSWORD, tn):
    
    global ZYXEL

    # 提示訊息
    print("\nHELLO ZYXEL!\n")

    # 輸入密碼
    tn.read_until(b"Password: ", timeout = 5)
    tn.write(PASSWORD.encode("ascii") + b"\n")

    # 如果登入失敗
    output = tn.read_until(b"Password ", timeout=5)
    if output == b'**********\r\n\r\nPassword: ' or output == b'*************\r\n\r\nPassword: ' or output == b'*****\r\n\r\nPassword: ':
        print("密碼錯誤. => 嘗試使用下一組密碼.")
        return 0

    # 輸入指令
    tn.write(b"show sys\r\n")
    tn.write(b"show ip arp\r\n")
    tn.write(b"show mac address-table all PORT\r\n")
    tn.write(b"show lldp info remote\r\n")
    time.sleep(3)
    output = tn.read_very_eager().decode("ascii")

    # 關閉連線
    tn.close()

    # 將telnet結果輸出至zyxel_output.txt檔案中
    with open("zyxel_output.txt", "w") as out:
        out.write(output)

    # 將ZYXEL TAG設成1
    ZYXEL = 1

    # fun5.輸出成功登入的密碼
    print_password(PASSWORD)
 
    #  回傳登入成功
    return 1

# fun4.登入edgecore
def log_in_edgecore(PASSWORD, tn):

    global ZYXEL

    # 提示訊息
    print("\nHELLO EDGECORE!\n")

    # 輸入密碼
    tn.read_until(b"Password: ", timeout = 5)
    tn.write(PASSWORD.encode("ascii") + b"\n")

    # 如果登入失敗
    output = tn.read_until(b"Username: ", timeout=5)
    if output == b'\r\nUsername: ':
        print("密碼錯誤. => 嘗試使用下一組密碼.")
        return 0

    # 輸入指令
    tn.write(b"terminal length 0\r\n")
    tn.write(b"show sys\r\n")
    tn.write(b"show arp\r\n")
    tn.write(b"show mac-address-table\r\n")
    tn.write(b"show lldp info remote-device\r\n")
    time.sleep(3)
    output = tn.read_very_eager().decode("ascii")

    # 關閉連線
    tn.close()

    # 將telnet結果輸出至edgecore_output.txt檔案中
    with open("edgecore_output.txt", "w") as out:
        out.write(output)

    # 將ZYXEL TAG設成0(edgecore)
    ZYXEL = 0

    # fun5.輸出成功登入的密碼
    print_password(PASSWORD)

    #  回傳登入成功
    return 1

# fun5.輸出成功登入的密碼
def print_password(PASSWORD):

    # 判斷是哪一組密碼並輸出
    if PASSWORD == "admin":
        print("Correct password : admin\n")
    elif PASSWORD == "1234":
        print("Correct password : 1234\n")
    else:
        print("Incorrect password.\n")

# fun6.整理EDGECORE輸出結果
def EDGECORE_parse(i):

    global mac_text

    # 讀檔
    with open("edgecore_output.txt", "r") as out:
        mac_text = out.read()

    # gateway_iparp
    make_gateway_iparp()

    # lldp_info資訊
    EDGECORE_make_lldp_info()

    # 查詢mac_vendor
    mac_vendor()

    # 輸出結果
    write_file(i)

# fun7.出口閘道iparp
def make_gateway_iparp():

    global pattern3, matches3

    # regex定義與比對
    if ZYXEL == 0:
        pattern3 = r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})\s+([0-9A-Fa-f-]+)\s+'
    else:
        pattern3 = r'\s+\d+\s+(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})\s+([0-9A-Fa-f:-]+)\s+'
    matches3  = re.findall(pattern3, mac_text)

    # 將比對結果輸出成gateway_iparp
    if trunk == "":
        with open("gateway_iparp.txt", "w") as out:
            for address in matches3:
                ip, mac = address
                out.write(ip)
                out.write(" ")
                
                mac = mac.upper()
                mac = mac.replace(":", "-")
                out.write(mac)
                out.write("\n")
    
# fun8.填入LLDP與比對IPARP結果
def EDGECORE_make_lldp_info():

    global lldp1, lldp2, lldp3, output_str2, dev_port, cnt

    # 切割出show lldp info remote-device後的字串
    lldp1 = mac_text.split("show lldp info remote-device")

    # 每一行切開
    lldp2 = lldp1[1].split("\n")
    
    # 遞迴每一行
    for cur in lldp2:

        # 用空格切開每一行的各個字串
        lldp3 = cur.split()

        # 利用LLDP有幾個字串判斷是否有加上備註
        if len(lldp3) == 4 and lldp3[0] == "Eth":

            # 排版
            if len(lldp3[1]) == 4:
                output_str2 = output_str2 + lldp3[0] + " " + lldp3[1] + "  " + lldp3[2] 
            else:
                output_str2 = output_str2 + lldp3[0] + " " + lldp3[1] + "   " + lldp3[2]

            output_str2 = output_str2 + "\n"

        elif len(lldp3) == 5:

            # 排版
            if len(lldp3[1]) == 4:
                output_str2 = output_str2 + lldp3[0] + " " + lldp3[1] + "  " + lldp3[2] + "  " + lldp3[4]
            else:
                output_str2 = output_str2 + lldp3[0] + " " + lldp3[1] + "   " + lldp3[2] + "  " + lldp3[4]

            output_str2 = output_str2 + "\n"

# fun9.查詢mac_vendor
def mac_vendor():

    global pattern1, matches1, output_str1

    # regex定義與比對
    if ZYXEL == 0:
        pattern1 = r'Eth \d+\/(\s*\d+)\s+([0-9A-Fa-f-]+)\s+\d+ Learn'
    else:
        pattern1 = r"\s+(\d+)\s+\d+\s+((?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2})\s+[A-Z]+"
    
    matches1  = re.findall(pattern1, mac_text)

    # 遞迴詢問
    for match in matches1:

        # 切割字串
        port = match[0]
        mac = match[1]
        url = url_prefix + mac + "/company/name"

        # 如果是上行port是trunk就不用往下處理
        if trunk != "" and int(port) == int(trunk):
            continue

        # API詢問
        if url.count('-') >= 5 or url.count(':') >= 5:
            response = requests.get(url)
            print(url, "\nMAC_VENDOR : ", response.text, "\n")
        else:
            continue

        # 如果http回應是200
        if response.status_code == 200:
            output_str1 = output_str1 + 'Port: ' + port + ', MAC: ' + mac + " ( " + response.text.strip() + " )\n"    
        else:
            if mac.count('-') == 5 or mac.count(':') == 5:
                output_str1 = output_str1 + 'Port: ' + port + ', MAC: ' + mac + '\n'
        
        # 間隔時間
        time.sleep(1)
    
# fun10.寫檔
def write_file(i):

    global pattern2, matches2, filename, pattern3, matches3, output_str2

    # regex定義與比對chassis mac-address
    if ZYXEL == 0:
        pattern2 = r'CPU \s+([0-9A-Fa-f-]+)\s+\d+ CPU'
    else:
        pattern2 = r'Ethernet Address\s+:\s+((?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2})'
    matches2 = re.findall(pattern2, mac_text)
    for j in range(len(matches2)):
        matches2[j] = matches2[j].upper()
        matches2[j] = matches2[j].replace(":", "-")

    # regex定義與比對設備型號
    if ZYXEL == 0:
        pattern3 = r'(ECS[0-9A-Za-z-]{7,})'
    else:
        pattern3 = r'([XGS]{2,3}+[0-9A-Za-z-]{7,})'
    matches3 = re.findall(pattern3, mac_text)
    for j in range(len(matches3)):
        matches3[j] = matches3[j].upper()
        matches3[j] = matches3[j].replace(":", "-")
    
    # 寫檔
    filename = ipList[i] + ".txt"
    with open(filename, "w") as out:
        out.write("=========chassis mac-address===========================================================\n\n")

    with open(filename, "a") as out:
        out.write("MAC : ")
        out.write(matches2[0])
        out.write("\nMODEL : ")
        out.write(matches3[0])
        out.write("\n")

    with open(filename, "a") as out:
        out.write("\n=========show mac-address-table===========================================================\n\n")

    with open(filename, "a") as out:
        out.write(output_str1)

    with open(filename, "a") as out:
        out.write("\n=========show lldp info remote-device===========================================================\n\n")

    with open(filename, "a") as out:
        out.write(output_str2)

# fun11.整理ZYXEL輸出結果
def ZYXEL_parse(i):

    global mac_text

    # 讀檔
    with open("zyxel_output.txt", "r") as out:
        mac_text = out.read()

    # gateway_iparp
    make_gateway_iparp()

    # lldp_info資訊
    ZYXEL_make_lldp_info()

    # 查詢mac_vendor
    mac_vendor()

    # 輸出結果
    write_file(i)

# fun12.填入LLDP資訊
def ZYXEL_make_lldp_info():
    global lldp1, lldp2, output_str2, pattern3

    # regex
    pattern3 = r"\s+(\d+)\s+((?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2})\s+"

    # 切割出LLDP INFO後的資料
    lldp1 = mac_text.split("LLDP Remote Device Information:")
    lldp2 = lldp1[1].split("\n")

    for cur in lldp2:

        matches3 = re.findall(pattern3, cur)

        if matches3 == []:
            continue
        else:
            s = str(matches3[0][1]).upper()
            s = s.replace(":", "-")

            # 排個版並輸出LLDP資訊
            if int(matches3[0][0]) < 10:
                output_str2 = output_str2 + "Port:" + " " + str(matches3[0][0]) + ", " + s
            else:
                output_str2 = output_str2 + "Port:" + str(matches3[0][0]) + ", " + s

        output_str2 = output_str2 + "\n"

# fun13.製作二階表裡的標頭
def make_doc_header(doc, i):
    global filename, doc_str1, pattern1, lines

    # initialize
    doc_str1 = ""
    pattern1 = r'MAC : ([0-9A-Fa-f-:]+)'

    # 新增一個段落並設定置中
    para1 = doc.add_paragraph()
    para1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = para1.add_run('關山國中骨幹交換器佈線點位表')
    run1.bold = True
    run1.font.size = Pt(20)

    # 新增設備階層與廠牌
    para2 = doc.add_paragraph()

    if ZYXEL == 0:
        run2 = para2.add_run('設備/階層:28交換器/第二層\t\t廠牌:  Edgecore')
    else:
        run2 = para2.add_run('設備/階層:28交換器/出口閘道\t廠牌:  ZYXEL')
    run2.bold = True
    run2.font.size = Pt(14)

    # 讀檔
    filename = ipList[i] + ".txt"
    with open(filename, "r") as f:
        lines = f.readlines()
    
    for line in lines:
        matches1 = re.findall(pattern1, line)
        if len(matches1) != 0:
            doc_str1 = matches1[0]

    # 新增設備位置與MAC
    para3 = doc.add_paragraph()
    run3 = para3.add_run('位置:\t\t\t\t\t\tMAC:  ' + doc_str1)
    run3.bold = True
    run3.font.size = Pt(14) 

# fun14.製作二階表裡的表格
def make_doc_table(doc, i):
    global pattern1, pattern2, pattern3, doc_str1

    # 開一個29列3行的表格
    table = doc.add_table(rows=29, cols=3, style='Table Grid')

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # fun15.幫表格上色
    shade_cells([table.cell(0, 0), table.cell(0, 1), table.cell(0, 2)], "#fce5cd")

    # 更換標頭文字
    header_cells = table.rows[0].cells
    header_cells[0].text = 'PORT'
    header_cells[0].width = Cm(2.09)
    header_cells[1].text = '對應目的地'
    header_cells[1].width = Cm(8.32)
    header_cells[2].text = '線路編號'
    header_cells[2].width = Cm(7.5)

    # word表格1-29列
    for j in range(1, 29):
        
        cells = table.rows[j].cells

        # 更改行寬 
        cells[0].width = Cm(2.09) 
        cells[1].width = Cm(8.32)
        cells[2].width = Cm(7.5)

        # 將每一行數字填入表格
        cells[0].text  = str(j)

        # 將設備MAC與VENDOR填入表格
        pattern1 = r'Port:\s+(\d+),'
        pattern2 = r'\((.*?)\)'
        pattern3 = r'MAC:\s+((?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2})'
        doc_str1 = ""
        for line in lines:
            matches1 = re.findall(pattern1, line)
            matches2 = re.findall(pattern2, line)
            matches3 = re.findall(pattern3, line)
            if len(matches1) != 0 and len(matches2) != 0 and len(matches3) != 0 and j == int(matches1[0]):
                print("%2d" % j, " : ", matches2, matches3)
                doc_str1 = doc_str1 + matches3[0] + "\n(" + matches2[0] + ")\n\n"
                print(matches2[0])
                cells[1].text = doc_str1

        # 如果是Trunk就填入上行
        if trunk != "" and j == int(trunk):
            cells[1].text = "上行"
        else:
            continue


    # 把表格的排列設定成垂直至中
    for row in table.rows:
        row.height = Cm(0.69)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 將標頭字體轉換成標楷體
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = '標楷體'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    # 將表格字體轉換成標楷體
    for row in table.rows:
        for cell in row.cells:
            para4 = cell.paragraphs

            for p in para4:
                for r in p.runs:
                    r.font.name = '標楷體'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

                
    # 編輯檔名 與 儲存檔案
    file_name = ipList[i] + ".docx"
    doc.save(file_name)

# fun15.幫表格上色
def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

#endregion

#region 主程式

# 分行讀取ip.txt裡的ip
with open("ip.txt", "r", encoding="utf-8") as f:
    ipList = f.readlines()

# 建立網路拓樸圖
G = nx.Graph()

# for迴圈設備數量
for i in range(len(ipList)):

    # initialize
    output_str1 = ""
    output_str2 = ""

    # fun1.處理字串
    if string_process(i) == 0:
        continue

    # 迴圈欲輸入的密碼
    for PASSWORD in PASSWORDS:
        # 如果密碼錯誤登入失敗 => 試下一組密碼(下一個迴圈)
        if log_in(PASSWORD) == 0:
            continue

        # 登入成功 => 跳出迴圈
        else:
            break            

    # 如果是EDGECORE設備
    if ZYXEL == 0:
        # fun6.整理EDGECORE輸出結果
        EDGECORE_parse(i)

    # 如果是ZYXEL設備
    else:
        # fun11.整理ZYXEL輸出結果
        ZYXEL_parse(i)

    # 建立新的word物件
    doc = Document()

    # fun13.製作二階表的標頭
    make_doc_header(doc, i)

    # fun14.製作二階表的表格
    make_doc_table(doc, i)

#endregion
        
